from pathlib import Path

from docx import Document
from pypdf import PdfReader, PdfWriter

from document_merger import merge_pdfs, merge_docx


def make_pdf(path: Path, pages: int):
    writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(width=300, height=300)
    with path.open("wb") as f:
        writer.write(f)


def make_docx(path: Path, text: str):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)


def test_merge_pdfs(tmp_path):
    a = tmp_path / "a.pdf"
    b = tmp_path / "b.pdf"
    out = tmp_path / "merged.pdf"
    make_pdf(a, 2)
    make_pdf(b, 3)

    merge_pdfs([a, b], out)

    assert len(PdfReader(str(out)).pages) == 5


def test_merge_docx(tmp_path):
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    out = tmp_path / "merged.docx"
    make_docx(a, "First document")
    make_docx(b, "Second document")

    merge_docx([a, b], out)

    result = Document(str(out))
    text = "\n".join(p.text for p in result.paragraphs)
    assert "First document" in text
    assert "Second document" in text
