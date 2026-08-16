from pathlib import Path

from docx import Document
from pypdf import PdfReader, PdfWriter

from document_tools import (
    extract_pages, merge_docx, merge_pdfs, reorder_pdf, rotate_pdf, split_pdf
)


def make_pdf(path: Path, pages: int):
    writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(width=300, height=300)
    with path.open("wb") as f:
        writer.write(f)


def make_docx(path: Path, text: str):
    d = Document()
    d.add_paragraph(text)
    d.save(path)


def test_merge_and_split(tmp_path):
    a, b, merged = tmp_path/"a.pdf", tmp_path/"b.pdf", tmp_path/"merged.pdf"
    make_pdf(a, 2); make_pdf(b, 3)
    merge_pdfs([a,b], merged)
    assert len(PdfReader(str(merged)).pages) == 5
    pages = split_pdf(merged, tmp_path/"split")
    assert len(pages) == 5


def test_extract_and_reorder(tmp_path):
    src, out1, out2 = tmp_path/"src.pdf", tmp_path/"extract.pdf", tmp_path/"order.pdf"
    make_pdf(src, 4)
    extract_pages(src, [1,3], out1)
    assert len(PdfReader(str(out1)).pages) == 2
    reorder_pdf(src, [4,2,1,3], out2)
    assert len(PdfReader(str(out2)).pages) == 4


def test_merge_docx(tmp_path):
    a, b, out = tmp_path/"a.docx", tmp_path/"b.docx", tmp_path/"merged.docx"
    make_docx(a, "First"); make_docx(b, "Second")
    merge_docx([a,b], out)
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "First" in text and "Second" in text
