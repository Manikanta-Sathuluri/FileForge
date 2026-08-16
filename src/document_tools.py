from __future__ import annotations

import shutil
import subprocess
from pathlib import Path
from typing import Callable, Iterable

from docx import Document
from docxcompose.composer import Composer
from pypdf import PdfReader, PdfWriter


def _progress(cb, done, total):
    if cb:
        cb(done, total)


def pdf_info(path: Path) -> dict:
    reader = PdfReader(str(path))
    return {
        "pages": len(reader.pages),
        "encrypted": bool(reader.is_encrypted),
        "size": path.stat().st_size,
    }


def merge_pdfs(paths: list[Path], output: Path, progress_callback=None) -> Path:
    if len(paths) < 2:
        raise ValueError("Select at least two PDF files.")
    writer = PdfWriter()
    for i, path in enumerate(paths, 1):
        reader = PdfReader(str(path))
        if reader.is_encrypted:
            try:
                result = reader.decrypt("")
            except Exception as exc:
                raise ValueError(f"'{path.name}' is password protected.") from exc
            if not result:
                raise ValueError(f"'{path.name}' is password protected.")
        for page in reader.pages:
            writer.add_page(page)
        _progress(progress_callback, i, len(paths))
    output.parent.mkdir(parents=True, exist_ok=True)
    with output.open("wb") as f:
        writer.write(f)
    return output


def split_pdf(path: Path, output_dir: Path, progress_callback=None) -> list[Path]:
    reader = PdfReader(str(path))
    if reader.is_encrypted:
        try:
            if not reader.decrypt(""):
                raise ValueError("The PDF is password protected.")
        except Exception as exc:
            raise ValueError("The PDF is password protected.") from exc

    output_dir.mkdir(parents=True, exist_ok=True)
    created = []
    total = len(reader.pages)
    stem = path.stem

    for i, page in enumerate(reader.pages, 1):
        out = output_dir / f"{stem}_page_{i:03d}.pdf"
        writer = PdfWriter()
        writer.add_page(page)
        with out.open("wb") as f:
            writer.write(f)
        created.append(out)
        _progress(progress_callback, i, total)
    return created


def extract_pages(path: Path, pages: list[int], output: Path) -> Path:
    reader = PdfReader(str(path))
    writer = PdfWriter()
    for page_number in pages:
        if page_number < 1 or page_number > len(reader.pages):
            raise ValueError(f"Page {page_number} is outside the document.")
        writer.add_page(reader.pages[page_number - 1])
    output.parent.mkdir(parents=True, exist_ok=True)
    with output.open("wb") as f:
        writer.write(f)
    return output


def rotate_pdf(path: Path, output: Path, degrees: int) -> Path:
    if degrees % 90 != 0:
        raise ValueError("Rotation must be 90, 180, or 270 degrees.")
    reader = PdfReader(str(path))
    writer = PdfWriter()
    for page in reader.pages:
        page.rotate(degrees)
        writer.add_page(page)
    output.parent.mkdir(parents=True, exist_ok=True)
    with output.open("wb") as f:
        writer.write(f)
    return output


def reorder_pdf(path: Path, order: list[int], output: Path) -> Path:
    reader = PdfReader(str(path))
    if sorted(order) != list(range(1, len(reader.pages) + 1)):
        raise ValueError("Page order must contain every page exactly once.")
    writer = PdfWriter()
    for n in order:
        writer.add_page(reader.pages[n - 1])
    output.parent.mkdir(parents=True, exist_ok=True)
    with output.open("wb") as f:
        writer.write(f)
    return output


def compress_pdf(path: Path, output: Path) -> Path:
    """Lossless structural optimization. It may not shrink every PDF."""
    reader = PdfReader(str(path))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)

    # pypdf cannot safely recompress arbitrary embedded images without
    # potentially changing quality. This keeps the operation lossless.
    try:
        writer.compress_identical_objects(remove_identicals=True, remove_orphans=True)
    except TypeError:
        try:
            writer.compress_identical_objects()
        except Exception:
            pass

    output.parent.mkdir(parents=True, exist_ok=True)
    with output.open("wb") as f:
        writer.write(f)
    return output


def extract_text_to_docx(path: Path, output: Path) -> Path:
    """PDF -> DOCX text extraction. Layout fidelity is intentionally limited."""
    reader = PdfReader(str(path))
    doc = Document()
    doc.add_heading(path.stem, level=1)

    for index, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        if index > 1:
            doc.add_page_break()
        doc.add_heading(f"Page {index}", level=2)
        if text.strip():
            for block in text.split("\n"):
                doc.add_paragraph(block)
        else:
            doc.add_paragraph("[No extractable text on this page.]")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    return output


def merge_docx(paths: list[Path], output: Path, progress_callback=None) -> Path:
    if len(paths) < 2:
        raise ValueError("Select at least two Word files.")
    master = Document(str(paths[0]))
    composer = Composer(master)
    _progress(progress_callback, 1, len(paths))
    for i, path in enumerate(paths[1:], 2):
        composer.append(Document(str(path)))
        _progress(progress_callback, i, len(paths))
    output.parent.mkdir(parents=True, exist_ok=True)
    composer.save(str(output))
    return output



def delete_pages(path: Path, pages_to_delete: list[int], output: Path) -> Path:
    reader = PdfReader(str(path))
    total = len(reader.pages)
    delete_set = set(pages_to_delete)
    if any(p < 1 or p > total for p in delete_set):
        raise ValueError("One or more page numbers are outside the document.")
    writer = PdfWriter()
    for index, page in enumerate(reader.pages, 1):
        if index not in delete_set:
            writer.add_page(page)
    if len(writer.pages) == 0:
        raise ValueError("You cannot delete every page from a PDF.")
    output.parent.mkdir(parents=True, exist_ok=True)
    with output.open("wb") as f:
        writer.write(f)
    return output


def pdf_to_images(path: Path, output_dir: Path, fmt: str = "PNG", dpi: int = 150):
    """Render PDF pages to images using PyMuPDF."""
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError(
            "PDF → Images requires PyMuPDF. Install it and try again."
        ) from exc

    output_dir.mkdir(parents=True, exist_ok=True)
    doc = fitz.open(str(path))
    created = []
    matrix = fitz.Matrix(dpi / 72.0, dpi / 72.0)

    for index, page in enumerate(doc, 1):
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        ext = "jpg" if fmt.upper() in {"JPG", "JPEG"} else "png"
        out = output_dir / f"{path.stem}_page_{index:03d}.{ext}"
        if ext == "jpg":
            pix.save(str(out), jpg_quality=92)
        else:
            pix.save(str(out))
        created.append(out)

    doc.close()
    return created

def find_libreoffice() -> str | None:
    candidates = [
        shutil.which("soffice"),
        shutil.which("libreoffice"),
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return str(candidate)
    return None


def docx_to_pdf(path: Path, output_dir: Path) -> Path:
    """Optional DOCX -> PDF through locally installed LibreOffice."""
    soffice = find_libreoffice()
    if not soffice:
        raise RuntimeError(
            "DOCX → PDF requires LibreOffice to be installed. "
            "Install LibreOffice, then try again."
        )
    output_dir.mkdir(parents=True, exist_ok=True)
    before = {p.resolve() for p in output_dir.glob("*.pdf")}
    cmd = [
        soffice, "--headless", "--convert-to", "pdf",
        "--outdir", str(output_dir), str(path)
    ]
    result = subprocess.run(
        cmd, capture_output=True, text=True, timeout=180,
        shell=False, check=False
    )
    if result.returncode != 0:
        raise RuntimeError((result.stderr or result.stdout).strip() or "LibreOffice conversion failed.")

    expected = output_dir / f"{path.stem}.pdf"
    if expected.exists():
        return expected

    new_files = [p for p in output_dir.glob("*.pdf") if p.resolve() not in before]
    if new_files:
        return new_files[0]
    raise RuntimeError("LibreOffice reported success but no PDF was produced.")
